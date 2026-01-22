"""
DOCX Generator for Dietician Agent Learning Guide
Run: pip install python-docx && python generate_docx.py
"""

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
import os

def create_learning_guide():
    doc = Document()
    
    # Title
    title = doc.add_heading('🧠 Dietician Agent - Complete Learning Guide', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    subtitle = doc.add_paragraph('A comprehensive guide to understanding how this AI-powered dietician application was built using Ollama, FastAPI, MongoDB, and React.')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    
    # Table of Contents
    doc.add_heading('Table of Contents', level=1)
    toc_items = [
        '1. Project Overview',
        '2. Architecture',
        '3. Technology Stack',
        '4. Backend Deep Dive',
        '5. Ollama Integration - Deep Dive',
        '6. Database Design',
        '7. API Endpoints',
        '8. Frontend Overview',
        '9. How to Explain This Project'
    ]
    for item in toc_items:
        doc.add_paragraph(item, style='List Number')
    
    # Section 1: Project Overview
    doc.add_page_break()
    doc.add_heading('1. Project Overview', level=1)
    
    doc.add_heading('What Does This Application Do?', level=2)
    doc.add_paragraph('The Dietician Agent is an AI-powered health assistant that:')
    overview_items = [
        'Analyzes lab reports using computer vision (OCR)',
        'Identifies health issues from extracted lab values',
        'Generates personalized 7-day meal plans based on health conditions',
        'Tracks patient progress over time to improve recommendations'
    ]
    for item in overview_items:
        doc.add_paragraph(item, style='List Bullet')
    
    doc.add_heading('Key Innovation', level=2)
    doc.add_paragraph('This project runs entirely on local hardware using Ollama - no cloud APIs, no subscription costs, and complete data privacy.')
    
    # Section 2: Architecture
    doc.add_page_break()
    doc.add_heading('2. System Architecture', level=1)
    
    doc.add_paragraph('The following diagram shows how all components interact:')
    
    # Add architecture image
    if os.path.exists('docs/system_architecture.png'):
        doc.add_picture('docs/system_architecture.png', width=Inches(6))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_heading('Data Flow', level=2)
    doc.add_paragraph('Here is the step-by-step flow of data through the system:')
    
    # Add data flow image
    if os.path.exists('docs/data_flow.png'):
        doc.add_picture('docs/data_flow.png', width=Inches(5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Section 3: Technology Stack
    doc.add_page_break()
    doc.add_heading('3. Technology Stack', level=1)
    
    # Technology table
    tech_table = doc.add_table(rows=7, cols=3)
    tech_table.style = 'Table Grid'
    
    # Header row
    header_cells = tech_table.rows[0].cells
    header_cells[0].text = 'Layer'
    header_cells[1].text = 'Technology'
    header_cells[2].text = 'Purpose'
    
    tech_data = [
        ('Frontend', 'React + Vite', 'User interface'),
        ('Backend', 'FastAPI (Python)', 'REST API server'),
        ('Database', 'MongoDB (Docker)', 'Store patients, reports, plans'),
        ('AI Runtime', 'Ollama', 'Run LLMs locally'),
        ('Vision Model', 'llava:7b', 'Extract text from images'),
        ('Text Model', 'mistral / llama3.2', 'Generate analysis & plans'),
    ]
    
    for i, (layer, tech, purpose) in enumerate(tech_data, 1):
        row_cells = tech_table.rows[i].cells
        row_cells[0].text = layer
        row_cells[1].text = tech
        row_cells[2].text = purpose
    
    doc.add_paragraph()
    doc.add_heading('Why These Technologies?', level=2)
    reasons = [
        ('FastAPI', 'Async Python framework, perfect for AI workloads'),
        ('MongoDB', 'Schema-less, handles varied lab report structures'),
        ('Ollama', 'Simple API for running open-source LLMs locally'),
        ('React', 'Modern, component-based UI'),
    ]
    for tech, reason in reasons:
        p = doc.add_paragraph()
        p.add_run(f'{tech}: ').bold = True
        p.add_run(reason)
    
    # Section 4: Backend Deep Dive
    doc.add_page_break()
    doc.add_heading('4. Backend Deep Dive', level=1)
    
    doc.add_heading('Project Structure', level=2)
    structure = """backend/
├── app/
│   ├── config.py           # Environment variables
│   ├── main.py             # FastAPI app setup
│   ├── database/
│   │   └── mongodb.py      # Database connection
│   ├── models/
│   │   ├── patient.py      # Pydantic models for patients
│   │   ├── lab_report.py   # Models for lab data
│   │   └── diet_plan.py    # Models for meal plans
│   ├── routes/
│   │   ├── patients.py     # Patient CRUD endpoints
│   │   ├── reports.py      # Lab report upload/analysis
│   │   └── diet.py         # Diet plan generation
│   └── services/
│       ├── ollama_client.py    # Wrapper for Ollama API
│       ├── ocr_service.py      # Lab report extraction
│       ├── diet_planner.py     # Meal plan generation
│       └── progress_analyzer.py # Track improvements
├── .env                    # Configuration
└── requirement.txt         # Dependencies"""
    
    doc.add_paragraph(structure).style = 'No Spacing'
    
    doc.add_heading('config.py - Configuration Management', level=2)
    config_code = """from pydantic_settings import BaseSettings

class Settings(BaseSettings):
    mongodb_url: str          # MongoDB connection string
    ollama_url: str           # Ollama server address
    vision_model: str         # Model for OCR (llava:7b)
    text_model: str           # Model for text (mistral)
    ollama_timeout: int       # Request timeout

    class Config:
        env_file = ".env"     # Load from .env file

settings = Settings()"""
    doc.add_paragraph(config_code).style = 'No Spacing'
    
    doc.add_paragraph()
    doc.add_paragraph('Why Pydantic Settings?')
    for item in ['Automatically reads from .env file', 'Type validation at startup', 'Easy to change models without code changes']:
        doc.add_paragraph(item, style='List Bullet')
    
    # Section 5: Ollama Integration
    doc.add_page_break()
    doc.add_heading('5. Ollama Integration - Deep Dive', level=1)
    
    doc.add_paragraph('This is the core innovation of the project!')
    
    doc.add_heading('What is Ollama?', level=2)
    doc.add_paragraph('Ollama is a tool that lets you run Large Language Models (LLMs) locally on your own computer. Think of it as "Docker for AI models."')
    
    # Comparison table
    compare_table = doc.add_table(rows=5, cols=3)
    compare_table.style = 'Table Grid'
    
    compare_data = [
        ('Feature', 'Cloud APIs (OpenAI)', 'Ollama (Local)'),
        ('Cost', '$0.01-0.06 per 1K tokens', 'Free forever'),
        ('Privacy', 'Data sent to cloud', 'Data stays local'),
        ('Speed', 'Depends on internet', 'Depends on GPU'),
        ('Availability', 'Requires internet', 'Works offline'),
    ]
    
    for i, row_data in enumerate(compare_data):
        row_cells = compare_table.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('How Ollama Works', level=2)
    
    if os.path.exists('docs/ollama_architecture.png'):
        doc.add_picture('docs/ollama_architecture.png', width=Inches(5.5))
        last_para = doc.paragraphs[-1]
        last_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run('Key Insight: ').bold = True
    p.add_run('Ollama is just an HTTP server. Your code doesn\'t need any special Ollama library - just make HTTP requests!')
    
    doc.add_heading('The ollama_client.py Code', level=2)
    
    ollama_code = """import httpx
import base64
from app.config import settings

class OllamaClient:
    def __init__(self):
        # Ollama server URL from .env
        self.base_url = settings.ollama_url
        self.timeout = httpx.Timeout(float(settings.ollama_timeout))
    
    async def generate(self, model: str, prompt: str, num_predict: int = 8192):
        \"\"\"Send a text prompt and get a response\"\"\"
        async with httpx.AsyncClient(timeout=self.timeout) as client:
            response = await client.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "stream": False,
                    "options": {
                        "num_predict": num_predict,
                        "num_ctx": 8192,
                    }
                },
            )
            return response.json()["response"]
    
    async def generate_with_image(self, model: str, prompt: str, image_path: str):
        \"\"\"Send an image to a vision model\"\"\"
        with open(image_path, "rb") as f:
            image_b64 = base64.b64encode(f.read()).decode()
        
        async with httpx.AsyncClient(timeout=self.timeout) as client:
            response = await client.post(
                f"{self.base_url}/api/generate",
                json={
                    "model": model,
                    "prompt": prompt,
                    "images": [image_b64],  # KEY: Base64 encoded images
                },
            )
            return response.json()["response"]"""
    
    doc.add_paragraph(ollama_code).style = 'No Spacing'
    
    doc.add_paragraph()
    doc.add_heading('Key Parameters', level=2)
    
    params_table = doc.add_table(rows=5, cols=2)
    params_table.style = 'Table Grid'
    
    params_data = [
        ('Parameter', 'Purpose'),
        ('num_predict', 'Maximum tokens to generate (increase for longer outputs)'),
        ('num_ctx', 'Context window size (input + output tokens)'),
        ('temperature', 'Randomness (0 = deterministic, 1 = creative)'),
        ('stream', 'Set false to get complete response at once'),
    ]
    
    for i, row_data in enumerate(params_data):
        row_cells = params_table.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('Common Ollama Issues & Solutions', level=2)
    
    issues_table = doc.add_table(rows=6, cols=3)
    issues_table.style = 'Table Grid'
    
    issues_data = [
        ('Issue', 'Cause', 'Solution'),
        ('Connection refused', 'Ollama not running', 'Run ollama serve on desktop'),
        ('Model not found', 'Model not downloaded', 'Run ollama pull model_name'),
        ('Response cut off', 'num_predict too low', 'Increase to 8192 or higher'),
        ('Timeout error', 'Model too slow', 'Increase OLLAMA_TIMEOUT in .env'),
        ('Out of memory', 'Model too big', 'Use smaller model (7b vs 13b)'),
    ]
    
    for i, row_data in enumerate(issues_data):
        row_cells = issues_table.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    # Section 6: Database Design
    doc.add_page_break()
    doc.add_heading('6. Database Design', level=1)
    
    doc.add_heading('Collections', level=2)
    
    doc.add_heading('patients', level=3)
    patients_schema = """{
    "_id": "ObjectId",
    "name": "John Doe",
    "email": "john@example.com",
    "dietary_preference": {
        "type": "veg",
        "allergies": ["nuts", "dairy"],
        "cuisines": ["Indian", "Mediterranean"],
        "meal_frequency": 3
    },
    "created_at": "2024-01-17T12:00:00Z"
}"""
    doc.add_paragraph(patients_schema).style = 'No Spacing'
    
    doc.add_paragraph()
    doc.add_heading('lab_reports', level=3)
    reports_schema = """{
    "_id": "ObjectId",
    "patient_id": "string",
    "extracted_data": {
        "hemoglobin": {"value": 14.5, "unit": "g/dL", "status": "normal"},
        "glucose": {"value": 126, "unit": "mg/dL", "status": "high"}
    },
    "health_analysis": {
        "issues": ["Elevated glucose indicates prediabetes"],
        "recommendations": ["Reduce sugar intake"]
    }
}"""
    doc.add_paragraph(reports_schema).style = 'No Spacing'
    
    # Section 7: API Endpoints
    doc.add_page_break()
    doc.add_heading('7. API Endpoints', level=1)
    
    doc.add_heading('Patients API', level=2)
    endpoints_table1 = doc.add_table(rows=5, cols=3)
    endpoints_table1.style = 'Table Grid'
    
    endpoints_data1 = [
        ('Method', 'Endpoint', 'Description'),
        ('POST', '/api/patients/', 'Register new patient'),
        ('GET', '/api/patients/search?email=...', 'Find patient by email'),
        ('GET', '/api/patients/{id}', 'Get patient details'),
        ('GET', '/api/patients/{id}/history', 'Get all reports and plans'),
    ]
    
    for i, row_data in enumerate(endpoints_data1):
        row_cells = endpoints_table1.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('Reports API', level=2)
    endpoints_table2 = doc.add_table(rows=3, cols=3)
    endpoints_table2.style = 'Table Grid'
    
    endpoints_data2 = [
        ('Method', 'Endpoint', 'Description'),
        ('POST', '/api/reports/upload', 'Upload and analyze lab report'),
        ('GET', '/api/reports/{id}', 'Get report details'),
    ]
    
    for i, row_data in enumerate(endpoints_data2):
        row_cells = endpoints_table2.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    doc.add_paragraph()
    doc.add_heading('Diet Plans API', level=2)
    endpoints_table3 = doc.add_table(rows=3, cols=3)
    endpoints_table3.style = 'Table Grid'
    
    endpoints_data3 = [
        ('Method', 'Endpoint', 'Description'),
        ('POST', '/api/diet/generate/{report_id}', 'Generate 7-day meal plan'),
        ('GET', '/api/diet/{id}', 'Get diet plan details'),
    ]
    
    for i, row_data in enumerate(endpoints_data3):
        row_cells = endpoints_table3.rows[i].cells
        for j, text in enumerate(row_data):
            row_cells[j].text = text
    
    # Section 8: Frontend Overview
    doc.add_page_break()
    doc.add_heading('8. Frontend Overview', level=1)
    
    doc.add_heading('Component Structure', level=2)
    frontend_structure = """frontend/src/
├── components/
│   ├── PatientForm.jsx      # New patient registration
│   ├── LabReportUpload.jsx  # Drag & drop file upload
│   ├── HealthAnalysis.jsx   # Display extracted lab data
│   └── DietPlanView.jsx     # 7-day meal plan display
├── services/
│   └── api.js               # API client for backend
├── App.jsx                  # Main app with navigation
└── index.css                # Dark theme styling"""
    doc.add_paragraph(frontend_structure).style = 'No Spacing'
    
    # Section 9: How to Explain
    doc.add_page_break()
    doc.add_heading('9. How to Explain This Project', level=1)
    
    doc.add_heading('Elevator Pitch (30 seconds)', level=2)
    pitch = doc.add_paragraph()
    pitch.add_run('"I built an AI-powered dietician that runs entirely on local hardware. It uses computer vision to read lab reports, identifies health issues like high cholesterol or diabetes risk, and generates personalized 7-day meal plans. The key innovation is using Ollama to run open-source models like Llava and Mistral locally, so there\'s no cloud dependency and complete data privacy."').italic = True
    
    doc.add_heading('Technical Deep Dive (5 minutes)', level=2)
    
    points = [
        ('1. The Problem:', 'Traditional health apps require expensive APIs like OpenAI or cloud-based services. I wanted to build something that runs completely locally.'),
        ('2. The Solution:', 'I used Ollama, which lets you run LLMs on consumer hardware. On my RTX 3070 with 8GB VRAM, I can run vision models and text generation locally.'),
        ('3. Key Challenges:', '• Vision models sometimes hallucinate values - I added validation\n• JSON parsing was tricky - LLMs don\'t always output valid JSON\n• Response truncation - had to increase num_predict to 8192 tokens'),
        ('4. What I Learned:', '• How to integrate AI models into production applications\n• The importance of prompt engineering for structured output\n• Async programming patterns with FastAPI\n• How to handle AI failures gracefully in UX'),
    ]
    
    for title, content in points:
        p = doc.add_paragraph()
        p.add_run(title).bold = True
        doc.add_paragraph(content)
    
    # Key Takeaways
    doc.add_page_break()
    doc.add_heading('Key Takeaways', level=1)
    
    takeaways = [
        'Ollama makes local AI accessible - No complex setup, just ollama pull model',
        'Prompt engineering is crucial - Getting structured JSON output requires careful prompts',
        'Error handling is non-negotiable - AI models can fail or produce garbage',
        'Environment variables are your friend - Easy to swap models without code changes',
        'MongoDB is great for AI apps - Schema-less design handles varied AI outputs',
    ]
    
    for i, takeaway in enumerate(takeaways, 1):
        doc.add_paragraph(f'{i}. {takeaway}')
    
    # Save
    doc.save('docs/LEARNING_GUIDE.docx')
    print('✅ Created: docs/LEARNING_GUIDE.docx')

if __name__ == '__main__':
    create_learning_guide()
